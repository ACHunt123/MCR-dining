#!/usr/bin/env python3
from docx import Document
import re,sys
'''
    Name Tag Formatter for MBS Event Attendees
    By A.C. Hunt | 04/13/2025

    This Python script writes the names of attendees to a name tags document, 
    taking the MBS name format as input. For example, the MBS names are listed as follows:   
        Name 1 \n
        Name 2 (n guests) \n
        etc.
    and are just copied and pasted from the MBS webpage into a text file.S
    
    The script processes the names and guest details as follows:
    
    1. It separates the main name and guests for each entry.
    2. Guests are formatted as 'Guest of Name X' and listed on a new line beneath the corresponding main name.
    3. The formatted list of names is inserted into a name tags document.

    If the MBS name file is 'unnamed', the script will generate a list of 35 unnamed guests:
        - 15 for Jesus College
        - 15 for Pembroke College
        - 5 for John's College

    These college assignments are hardcoded in the script.

    Note: The script currently doesn't handle situations where there are more names than the available name cards in the document, though this could be a potential feature for future versions.

    Usage Instructions:
    --------------------
    1. Copy and paste the MBS names from the webpage into a text file (e.g., "mbs_names.txt").
    2. Ensure you have a set of blank name tags in the same directory, with the placeholder 'Guest Name' where the name should appear.
    3. Run the script with the following arguments:
       - The name of the name tags document (e.g., 'name_tags_template.docx')
       - The desired output document name (e.g., 'formatted_name_tags.docx')
       - The name of the MBS names file (e.g., 'mbs_names.txt')

    Example command to run the script:
    python3 name_tag_formatter.py
    [the names are hardcoded, so just change them in the script]
'''


def inputnames(docx_file, outname, names_file):
    # Load the .docx file for editing
    doc = Document(docx_file)
    if names_file == 'unnamed': #  unnamed guests 
        attendees = []
        for i in range(15):
            attendees.append("Jesus College Guest\n")
        for i in range(15):
            attendees.append("Pembroke College Guest\n")
        for i in range(5):
            attendees.append("John's College Guest\n")
    else: # Load the list of guest names
        file = open(names_file, "r")
        lines = file.readlines()
        file.close()

        # unpack number of guests if listed in brackets
        attendees = []
        for line in lines:
            if 'MCR Guest' in line: continue # skip MCR guests
            # get any parentehsis (as these mean there are guests)
            if "(" in line:
                name = line.split("(")[0]                   # get the name (without the number of guests in brackets)
                attendees.append(f'{name}\n')               # add the name to the list
                numguests = re.findall(r'\d+', line)[0]     # get the number of guests
                for i in range(int(numguests)):
                    attendees.append(f'Guest of {name}\n')  # Add guest of x, numguests times
            else:
                attendees.append(line)                       # add the name to the list

    # center the names
    for i in range(len(attendees)):
        attendees[i] = attendees[i].strip() + "\n"


    i=0 # Counter for iterating through the attendees list
    len_names = len(attendees) # Number of names in the attendees list


    # Insert names into the document
    # Document is formatted as a table (see printdocstructure.py for more details)
    table=doc.tables[0] 
    for row_idx, row in enumerate(table.rows):
        cell = row.cells[0]
        for paragraph in cell.paragraphs:
            # Iterate through each run in the paragraph
            for run in paragraph.runs:
                # If "guest name" is found in the run text, replace it
                if "Guest Name" in run.text:  
                    if i>=len_names:
                        # remove the  paragraph if there are no more names
                        p = paragraph._element
                        p.getparent().remove(p)
                        continue
                    # Replace 'guest name' with the current name in the attendees list
                    run.text = run.text.replace("Guest Name", attendees[i])
                    i+=1
    
    if i<len_names: 
        print("Not all names were used, you will need to add more cards")

    # Save the modified document
    doc.save(outname)
    print(f"Saved to {outname}")
    print("------")

### Main
# Generate name tags for named attendees
inputnames("superhall_nametags.docx", "tags_superhall_named.docx", "superhall_names.txt")
# Generate name tags for unnamed attendees
inputnames("superhall_nametags.docx", "tags_superhall_unnamed.docx", "unnamed")

print("All done!")
